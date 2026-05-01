#!/usr/bin/env python3
"""
Consolida o log de aprendizado SavyCode (file-first) em arquivo .docx.

Modelo file-first:
    O log markdown contém apenas a explicação por construção, com referências
    de linha apontando ao arquivo real:

        ## Construção: `def` (L1 em `soma.py`)
        - **O que é:** ...

    Este script resolve cada referência, lê o trecho correspondente no
    arquivo real e injeta o código no .docx final ao lado da explicação.

Uso:
    python3 consolidate_to_docx.py \\
        --log <caminho/do/log.md> \\
        --project <caminho/do/projeto> \\
        [--profile ~/.claude/savycode/knowledge-profile.json] \\
        [--context-lines 0]

O arquivo gerado vai em: <projeto>/learning-journal/<basename-do-log>.docx

Dependência: python-docx (pip install python-docx). Se ausente, gera fallback .md formatado.
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from datetime import datetime, timezone
from pathlib import Path

# `## Construção: `nome` (L<linha>[-L<linha>] em `arquivo`)`
CONSTRUCT_RE = re.compile(
    r"^##\s+Construção:\s+`?(?P<name>[^`\n]+?)`?"
    r"\s*\(L(?P<line_start>\d+)(?:-L(?P<line_end>\d+))?"
    r"\s+em\s+`(?P<file>[^`]+)`\)\s*\n"
    r"(?P<body>(?:- .+\n?)+)",
    re.MULTILINE,
)

LESSON_RE = re.compile(r"^##\s+Lição:\s+(.+)$", re.MULTILINE)
TITLE_RE = re.compile(r"^#\s+Sessão:\s+(.+)$", re.MULTILINE)


def load_profile(profile_path: Path) -> dict:
    default = {"level": 1, "topics": {}, "weak_areas": []}
    if not profile_path.exists():
        return default
    try:
        raw = profile_path.read_text(encoding="utf-8").strip()
        return json.loads(raw) if raw else default
    except (json.JSONDecodeError, OSError):
        return default


def read_snippet(project_root: Path, rel_file: str, line_start: int, line_end: int | None, ctx: int) -> tuple[str, str]:
    """Retorna (linguagem-do-arquivo, snippet-do-código).

    Resolve o caminho relativamente à raiz do projeto. Se o arquivo não existir,
    devolve marcador `(arquivo não encontrado)` para preservar a entrega.
    """
    target = (project_root / rel_file).resolve()
    try:
        target.relative_to(project_root.resolve())
    except ValueError:
        return ("text", f"(referência inválida: {rel_file} fora da raiz do projeto)")

    if not target.exists():
        return ("text", f"(arquivo não encontrado: {rel_file})")

    lines = target.read_text(encoding="utf-8", errors="replace").splitlines()
    end = line_end or line_start
    lo = max(1, line_start - ctx)
    hi = min(len(lines), end + ctx)
    snippet = "\n".join(lines[lo - 1:hi])
    return (target.suffix.lstrip(".") or "text", snippet)


def parse_log(log_path: Path, project_root: Path, ctx: int) -> dict:
    text = log_path.read_text(encoding="utf-8")

    title_match = TITLE_RE.search(text)
    title = title_match.group(1).strip() if title_match else "Sessão sem título"

    constructs = []
    for m in CONSTRUCT_RE.finditer(text):
        ls = int(m.group("line_start"))
        le = int(m.group("line_end")) if m.group("line_end") else None
        lang, snippet = read_snippet(project_root, m.group("file"), ls, le, ctx)
        constructs.append({
            "name": m.group("name").strip(),
            "file": m.group("file"),
            "line_start": ls,
            "line_end": le,
            "body": m.group("body").strip(),
            "lang": lang,
            "snippet": snippet,
        })

    lessons = [m.group(1).strip() for m in LESSON_RE.finditer(text)]

    return {"title": title, "constructs": constructs, "lessons": lessons}


def write_docx(out_path: Path, sections: dict, profile: dict) -> bool:
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt
    except ImportError:
        return False

    doc = Document()
    doc.add_heading(f"Aprendizado: {sections['title']}", level=0)

    meta = doc.add_paragraph()
    meta.add_run(f"Gerado em: {datetime.now(timezone.utc).strftime('%Y-%m-%dT%H:%M:%SZ')}").italic = True
    meta.add_run(f"\nNível atual: {profile.get('level', 1)}").italic = True

    doc.add_heading("1. Contexto / Sessão", level=1)
    doc.add_paragraph(sections["title"])

    if sections["constructs"]:
        doc.add_heading("2. Construções e código real", level=1)
        for c in sections["constructs"]:
            range_label = f"L{c['line_start']}" + (f"-L{c['line_end']}" if c["line_end"] else "")
            doc.add_heading(f"{c['name']}  —  {c['file']} ({range_label})", level=2)
            doc.add_paragraph(c["body"])
            p = doc.add_paragraph()
            run = p.add_run(c["snippet"])
            run.font.name = "Menlo"
            run.font.size = Pt(9)

    if sections["lessons"]:
        doc.add_heading("3. Lições principais", level=1)
        for lesson in sections["lessons"]:
            doc.add_paragraph(lesson, style="List Bullet")

    doc.add_heading("4. Próximos estudos sugeridos", level=1)
    weak = profile.get("weak_areas", [])
    if weak:
        for area in weak:
            doc.add_paragraph(area, style="List Bullet")
    else:
        doc.add_paragraph("Nenhuma área fraca destacada no perfil atual.")

    doc.save(str(out_path))
    return True


def write_markdown_fallback(out_path: Path, sections: dict, profile: dict) -> Path:
    # Sufixo `.consolidated.md` evita colisão com o log de origem (mesmo stem).
    md_path = out_path.with_suffix(".consolidated.md")
    lines = [
        f"# Aprendizado: {sections['title']}",
        "",
        f"_Gerado em: {datetime.now(timezone.utc).strftime('%Y-%m-%dT%H:%M:%SZ')}_",
        f"_Nível atual: {profile.get('level', 1)}_",
        "",
        "## 1. Contexto / Sessão",
        sections["title"],
        "",
    ]
    if sections["constructs"]:
        lines.append("## 2. Construções e código real")
        for c in sections["constructs"]:
            range_label = f"L{c['line_start']}" + (f"-L{c['line_end']}" if c["line_end"] else "")
            lines.append(f"### {c['name']} — `{c['file']}` ({range_label})")
            lines.append(c["body"])
            lines.append(f"```{c['lang']}")
            lines.append(c["snippet"])
            lines.append("```")
            lines.append("")
    if sections["lessons"]:
        lines.append("## 3. Lições principais")
        for lesson in sections["lessons"]:
            lines.append(f"- {lesson}")
        lines.append("")
    lines.append("## 4. Próximos estudos sugeridos")
    for area in profile.get("weak_areas", []) or ["Nenhuma área fraca destacada."]:
        lines.append(f"- {area}")
    md_path.write_text("\n".join(lines), encoding="utf-8")
    return md_path


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("--log", required=True, type=Path, help="Caminho do log .md da sessão")
    parser.add_argument("--project", required=True, type=Path, help="Raiz do projeto (resolve referências Lxx em arquivo)")
    parser.add_argument(
        "--profile",
        type=Path,
        default=Path.home() / ".claude" / "savycode" / "knowledge-profile.json",
        help="Caminho do knowledge-profile.json",
    )
    parser.add_argument(
        "--context-lines",
        type=int,
        default=0,
        help="Linhas de contexto antes/depois do trecho referenciado (default: 0)",
    )
    args = parser.parse_args()

    if not args.log.exists():
        print(f"Log não encontrado: {args.log}", file=sys.stderr)
        return 1
    if not args.project.exists():
        print(f"Raiz do projeto não encontrada: {args.project}", file=sys.stderr)
        return 1

    target_dir = args.project / "learning-journal"
    target_dir.mkdir(parents=True, exist_ok=True)
    out_path = target_dir / (args.log.stem + ".docx")

    sections = parse_log(args.log, args.project, args.context_lines)
    profile = load_profile(args.profile)

    if not sections["constructs"]:
        print(
            "Aviso: nenhuma construção parseada. Verifique se o log usa o formato"
            " `## Construção: `nome` (Lxx em `arquivo`)`.",
            file=sys.stderr,
        )

    if write_docx(out_path, sections, profile):
        print(f"OK: {out_path}")
    else:
        md_path = write_markdown_fallback(out_path, sections, profile)
        print(f"python-docx não instalado. Gerado fallback markdown: {md_path}", file=sys.stderr)

    return 0


if __name__ == "__main__":
    sys.exit(main())
